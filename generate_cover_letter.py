import anthropic
import os
import json
from docx import Document
from config import CV_PATH, JD_PATH, OUTPUTS_DIR


def extract_job_details(jd_text, client):
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=200,
        messages=[{
            "role": "user",
            "content": f"""Extract the job title and company name from this job description.
Return only JSON in exactly this format, no other text:
{{
    "job_title": "the job title",
    "company_name": "the company name"
}}

JOB DESCRIPTION:
{jd_text}"""
        }]
    )

    response_text = message.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
    return json.loads(response_text.strip())


def generate_cover_letter(jd_text, cv_text, job_title, company_name, client):
    with open("cover_letter_prompt.txt", "r") as f:
        prompt_template = f.read()

    prompt = prompt_template.format(
        jd_text=jd_text,
        cv_text=cv_text,
        job_title=job_title,
        company_name=company_name
    )

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}]
    )

    return message.content[0].text.strip()


def build_cover_letter_docx(cover_letter_text, job_title, company_name, filename):
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    name_para = doc.add_paragraph()
    name_run = name_para.add_run("Subhash Chandra Yadav")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    name_run.font.name = "Calibri"

    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run(
        "+45 9180 7190  ·  syadav@allthingsagile.net  ·  linkedin.com/in/subhashchandrayadav  ·  Copenhagen, Denmark"
    )
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    contact_run.font.name = "Calibri"

    divider = doc.add_paragraph()
    divider_run = divider.add_run("─" * 80)
    divider_run.font.size = Pt(8)
    divider_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    divider_run.font.name = "Calibri"
    divider.paragraph_format.space_after = Pt(12)

    role_para = doc.add_paragraph()
    role_run = role_para.add_run(f"Re: {job_title} — {company_name}")
    role_run.bold = True
    role_run.font.size = Pt(11)
    role_run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
    role_run.font.name = "Calibri"
    role_para.paragraph_format.space_after = Pt(16)

    paragraphs = [p.strip() for p in cover_letter_text.split('\n\n') if p.strip()]

    for para_text in paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(para_text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = Pt(14)

    signoff = doc.add_paragraph()
    signoff_run = signoff.add_run("Kind regards,")
    signoff_run.font.size = Pt(10.5)
    signoff_run.font.name = "Calibri"
    signoff_run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)

    name_sign = doc.add_paragraph()
    name_sign_run = name_sign.add_run("Subhash Chandra Yadav")
    name_sign_run.bold = True
    name_sign_run.font.size = Pt(10.5)
    name_sign_run.font.name = "Calibri"
    name_sign_run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)

    doc.save(filename)


def main(job_title=None, company_name=None, output_dir=None):
    print("\nReading job description...")
    with open(JD_PATH, "r") as f:
        jd_text = f.read()

    print("Reading CV...")
    doc = Document(CV_PATH)
    cv_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    if not job_title or not company_name:
        print("Extracting job details...")
        job_details = extract_job_details(jd_text, client)
        job_title = job_details["job_title"]
        company_name = job_details["company_name"]
        print(f"  → {job_title} at {company_name}")

    if output_dir is None:
        folder_name = f"{job_title} @ {company_name}".replace("/", "-")
        output_dir = os.path.join(OUTPUTS_DIR, folder_name)
    os.makedirs(output_dir, exist_ok=True)

    print("Generating cover letter...")
    cover_letter_text = generate_cover_letter(
        jd_text, cv_text, job_title, company_name, client
    )

    filename = os.path.join(
        output_dir,
        f"Subhash_Yadav_Cover_Letter_{job_title}_{company_name}.docx".replace(" ", "_")
    )
    build_cover_letter_docx(cover_letter_text, job_title, company_name, filename)

    print(f"\nDone! Cover letter saved to: {filename}")
    return job_title, company_name, output_dir


if __name__ == "__main__":
    main()